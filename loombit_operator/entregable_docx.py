"""
entregable_docx.py — variante .docx del entregable autónomo (Skill W Administration Core).

Los gestores viven en Word: además del dossier HTML (`entregable.py`), aquí se renderiza el mismo
Expediente a un `.docx` editable. **python-docx es dependencia OPCIONAL**: se importa de forma
perezosa y, si falta, se levanta `DocxNoDisponibleError` (el router lo traduce a 501). Así el núcleo
no depende de ella y los entornos mínimos (p.ej. edge) siguen arrancando.

Determinista (lo construye CÓDIGO, no el LLM) e idéntico en contenido al dossier HTML, incluido el
**sello de integridad** (`verify_chain`). Local-first: no sale ni un byte a la red.
"""

from __future__ import annotations

import importlib.util
import io
import json
from pathlib import Path
from typing import Any

from .expedientes import Document, Expediente, ExpedienteEvent, ExpedienteStore

DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None


class DocxNoDisponibleError(RuntimeError):
    """python-docx no está instalado: el export .docx no está disponible."""


def _fmt_value(value: Any) -> str:
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    if value is None:
        return "—"
    return str(value)


def render_dossier_docx(
    exp: Expediente,
    events: list[ExpedienteEvent],
    documents: list[Document],
    *,
    chain_ok: bool,
    generated_at: str | None = None,
) -> bytes:
    """Renderiza el dossier como documento Word (.docx) y devuelve sus bytes."""
    if not DOCX_AVAILABLE:
        raise DocxNoDisponibleError("python-docx no está instalado; no se puede exportar a .docx")
    from datetime import UTC, datetime

    from docx import Document as DocxDocument  # import perezoso (dep opcional)

    generated_at = generated_at or datetime.now(UTC).isoformat()
    doc = DocxDocument()

    doc.add_heading(exp.title, level=0)
    doc.add_paragraph(f"{exp.kind} · {exp.status.value} · entidad: {exp.entity_id}")
    doc.add_paragraph(f"Expediente {exp.id}")
    doc.add_paragraph(f"Creado: {exp.created_at} · Actualizado: {exp.updated_at}")

    doc.add_heading("Resumen", level=1)
    if exp.data:
        tabla = doc.add_table(rows=0, cols=2)
        tabla.style = "Table Grid"
        for clave, valor in exp.data.items():
            celdas = tabla.add_row().cells
            celdas[0].text = str(clave)
            celdas[1].text = _fmt_value(valor)
    else:
        doc.add_paragraph("Sin datos registrados.")

    doc.add_heading("Documentos", level=1)
    if documents:
        for d in documents:
            doc.add_paragraph(
                f"{d.kind}: {Path(d.path).name} — sha256 {d.sha256[:16]}… ({d.added_at})",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("Sin documentos adjuntos.")

    doc.add_heading("Trazabilidad", level=1)
    if events:
        for ev in events:
            detalle = json.dumps(ev.detail, ensure_ascii=False) if ev.detail else ""
            doc.add_paragraph(
                f"{ev.kind} — {ev.ts} · {ev.actor} · {ev.hash[:12]}…  {detalle}".rstrip(),
                style="List Number",
            )
    else:
        doc.add_paragraph("Sin eventos.")

    doc.add_heading("Sello de integridad", level=1)
    if chain_ok:
        doc.add_paragraph(
            "✔ Integridad verificada: la cadena de hashes es consistente; el contenido no se ha "
            "alterado."
        )
    else:
        doc.add_paragraph(
            "✗ Integridad NO verificada: la cadena de hashes no cuadra; el expediente pudo "
            "manipularse. No lo tomes como prueba sin revisar el original."
        )
    doc.add_paragraph(f"Copia generada el {generated_at}. Documento producido por Loombit (local).")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_dossier_docx(store: ExpedienteStore, expediente_id: str) -> bytes:
    """Reúne expediente + eventos + documentos + verificación de cadena y devuelve el .docx."""
    exp = store.get(expediente_id)  # lanza ExpedienteNotFoundError si no existe
    return render_dossier_docx(
        exp,
        store.events(expediente_id),
        store.documents(expediente_id),
        chain_ok=store.verify_chain(expediente_id),
    )
